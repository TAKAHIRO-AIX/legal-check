import io
import json
import os
import sqlite3
import tempfile
import time
from datetime import datetime

import anthropic
import boto3
import httpx
import openpyxl
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.requests import Request

app = FastAPI(title="契約書リーガルチェックシステム")
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

AWS_REGION = os.getenv("AWS_REGION", "us-east-1")
ANTHROPIC_MODEL = "claude-sonnet-4-6"

textract = boto3.client("textract", region_name=AWS_REGION)
claude_client = anthropic.Anthropic()

# --- DB初期化 ---
DB_PATH = os.path.join(os.path.dirname(__file__), "history.db")

def init_db():
    conn = sqlite3.connect(DB_PATH)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS checks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT,
            created_at TEXT,
            background TEXT,
            focus TEXT,
            usage_method TEXT,
            output_indicator TEXT,
            other_context TEXT,
            primary_docs TEXT,
            secondary_docs TEXT,
            result TEXT,
            word_path TEXT
        )
    """)
    conn.commit()
    conn.close()

init_db()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """WordファイルからテキストをExtract"""
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """ExcelファイルをシートごとにチャンクとしてExtract"""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    chunks = []
    for sheet in wb.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            row_text = "\t".join(str(c) if c is not None else "" for c in row)
            if row_text.strip():
                rows.append(row_text)
        if rows:
            chunks.append(f"--- シート: {sheet.title} ---\n" + "\n".join(rows))
    return "\n\n".join(chunks)


def extract_text_from_txt(file_bytes: bytes) -> str:
    """テキストファイルをExtract"""
    return file_bytes.decode("utf-8", errors="replace")


def fetch_text_from_url(url: str) -> str:
    """URLからWebページのテキストを取得"""
    try:
        resp = httpx.get(url, timeout=30, follow_redirects=True, headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        # スクリプト・スタイルタグを除去
        for tag in soup(["script", "style", "nav", "header", "footer"]):
            tag.decompose()
        text = soup.get_text(separator="\n", strip=True)
        # 空行を圧縮
        lines = [line for line in text.split("\n") if line.strip()]
        return "\n".join(lines)
    except Exception as e:
        return f"[URL取得エラー: {str(e)}]"


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """TextractでスキャンPDFからテキスト抽出（非同期ジョブ）"""
    s3 = boto3.client("s3", region_name=AWS_REGION)
    bucket = os.getenv("TEXTRACT_S3_BUCKET")

    if not bucket:
        raise HTTPException(status_code=500, detail="環境変数 TEXTRACT_S3_BUCKET が未設定です")

    key = f"legal-check-tmp/{int(time.time())}.pdf"
    s3.put_object(Bucket=bucket, Key=key, Body=pdf_bytes)

    try:
        response = textract.start_document_text_detection(
            DocumentLocation={"S3Object": {"Bucket": bucket, "Name": key}}
        )
        job_id = response["JobId"]

        # ジョブ完了待ち
        while True:
            result = textract.get_document_text_detection(JobId=job_id)
            status = result["JobStatus"]
            if status == "SUCCEEDED":
                break
            if status == "FAILED":
                raise HTTPException(status_code=500, detail="Textractのテキスト抽出に失敗しました")
            time.sleep(3)

        # 全ページのテキストを結合
        pages: dict[int, list[str]] = {}
        next_token = None
        while True:
            kwargs = {"JobId": job_id}
            if next_token:
                kwargs["NextToken"] = next_token
            result = textract.get_document_text_detection(**kwargs)
            for block in result.get("Blocks", []):
                if block["BlockType"] == "LINE":
                    page = block.get("Page", 1)
                    pages.setdefault(page, []).append(block["Text"])
            next_token = result.get("NextToken")
            if not next_token:
                break

        lines = []
        for page_num in sorted(pages.keys()):
            lines.append(f"\n--- ページ {page_num} ---")
            lines.extend(pages[page_num])
        return "\n".join(lines)
    finally:
        s3.delete_object(Bucket=bucket, Key=key)


def legal_check_with_claude_v2(primary_docs: list[dict], secondary_docs: list[dict], background: str, focus: str, usage_method: str = "", output_indicator: str = "", other_context: str = "") -> list[dict]:
    """複数インプット対応のリーガルチェック"""
    
    # トークン上限対策: 各テキストを最大80,000文字に制限（合計で約800K tokens以内に収める）
    MAX_CHARS_PER_DOC = 80000
    for doc in primary_docs:
        if len(doc["text"]) > MAX_CHARS_PER_DOC:
            doc["text"] = doc["text"][:MAX_CHARS_PER_DOC] + "\n\n[※ テキストが長すぎるため、ここで切り詰めています]"
    for doc in secondary_docs:
        if len(doc["text"]) > MAX_CHARS_PER_DOC:
            doc["text"] = doc["text"][:MAX_CHARS_PER_DOC] + "\n\n[※ テキストが長すぎるため、ここで切り詰めています]"
    
    # チェック対象テキストを構築
    primary_section = ""
    has_ocr = any(d.get("is_ocr") for d in primary_docs)
    for i, doc in enumerate(primary_docs, 1):
        primary_section += f"\n### チェック対象 {i}: {doc['desc']}\n"
        if doc.get("is_ocr"):
            primary_section += "（※ スキャンPDFからOCR読取。文字化け・誤認識の可能性あり）\n"
        primary_section += doc["text"] + "\n"
    
    # 参照資料テキストを構築
    secondary_section = ""
    if secondary_docs:
        secondary_section = "\n## 参照資料（比較・参考用。これ自体はチェック対象ではない）\n"
        for i, doc in enumerate(secondary_docs, 1):
            secondary_section += f"\n### 参照資料 {i}: {doc['desc']}\n"
            secondary_section += doc["text"] + "\n"
    
    # OCR注意書き
    ocr_note = ""
    if has_ocr:
        ocr_note = """
- 一部のチェック対象はスキャンPDFからOCR読取したテキストです。文字化け・誤認識箇所がある可能性があります
- 判読不能な箇所は推測で補わず、その旨を指摘してください（ただし過度に強調せず、読める部分の法的分析を優先）"""

    # コンテキスト③④⑤を構築
    additional_context = ""
    if usage_method:
        additional_context += f"\n## 情報の活用方法（利用者からの指示）\n{usage_method}\n"
    if other_context:
        additional_context += f"\n## その他の指示・補足\n{other_context}\n"
    
    # 出力形式の指示
    output_instruction = """## 出力形式
結果は必ず以下のJSON配列形式のみで返してください（説明文不要、コードブロック不要）：
[
  {
    "page": ページ番号(int、不明なら1),
    "clause": "条項番号または条項名",
    "risk_level": "高/中/低",
    "issue": "問題点の説明",
    "recommendation": "修正・対応の提案"
  }
]"""
    if output_indicator:
        output_instruction = f"""## 出力条件・出力基準（利用者指定）
{output_indicator}

## 出力形式（システム指定・必須）
上記の出力条件を踏まえつつ、結果は必ず以下のJSON配列形式のみで返してください（説明文不要、コードブロック不要）：
[
  {{
    "page": ページ番号(int、不明なら1),
    "clause": "条項番号または条項名",
    "risk_level": "高/中/低",
    "issue": "問題点の説明",
    "recommendation": "修正・対応の提案"
  }}
]"""

    prompt = f"""あなたは日本法に精通した法律の専門家です。以下の情報に基づきリーガルチェックを実施してください。

## 重要な前提
- あなたの分析対象は「チェック対象」セクションの契約書です
- 「参照資料」は比較・参考のために提供されています。参照資料自体のチェックは不要です
- 「コンテキスト」は利用者の意図・希望を示す参考情報です
- 契約書の言語が日本語以外（韓国語・英語等）を含む場合、それぞれの言語部分を区別して分析してください{ocr_note}

## コンテキスト

### ① 背景・経緯・概要
{background}

### ② 注目点・着眼点・希望
{focus}
{additional_context}

## チェック対象の情報（STEP1）
{primary_section}
{secondary_section}

## 指示
上記のチェック対象を精査し、以下の観点でリスクや問題点を洗い出してください：
1. 利用者に不利な条項
2. 必須条項の欠落（秘密保持・準拠法・管轄裁判所・損害賠償上限・契約期間・解除条件など）
3. 曖昧・不明確な表現
4. 利用者の注目点に関連する事項
5. 参照資料（標準契約書等）と比較して不足・乖離している点

重要：各指摘は簡潔に（issue・recommendationは各200文字以内）。重要度の高いものから優先してください。利用者が出力条件で件数を指定した場合はそちらに従ってください。指定がない場合は全件出力してください。

{output_instruction}"""

    response = claude_client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=64000,
        messages=[{"role": "user", "content": prompt}],
    )
    result_text = response.content[0].text

    # JSONを抽出
    import re
    # ```json ... ``` ブロックを探す
    json_match = re.search(r'```(?:json)?\s*(\[[\s\S]*?\])\s*```', result_text)
    if json_match:
        try:
            return json.loads(json_match.group(1))
        except json.JSONDecodeError:
            pass
    
    # 直接 [ ... ] を探す
    start = result_text.find("[")
    end = result_text.rfind("]")
    if start >= 0 and end > start:
        json_str = result_text[start:end + 1]
        try:
            return json.loads(json_str)
        except json.JSONDecodeError:
            last_brace = json_str.rfind("}")
            if last_brace > 0:
                json_str = json_str[:last_brace + 1] + "]"
                try:
                    return json.loads(json_str)
                except json.JSONDecodeError:
                    pass
    
    raise ValueError(f"JSONが見つかりません: {result_text[:300]}")


def generate_word_report(issues: list[dict], background: str, focus: str) -> str:
    """python-docxでWordレポート生成、一時ファイルパスを返す"""
    doc = Document()

    # タイトル
    title = doc.add_heading("契約書リーガルチェックレポート", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    # コンテキスト
    doc.add_heading("■ チェック背景・コンテキスト", level=2)
    doc.add_paragraph(f"【背景・経緯】\n{background}")
    doc.add_paragraph(f"【注目点・着眼点】\n{focus}")

    # サマリー
    doc.add_heading("■ リスクサマリー", level=2)
    high = sum(1 for i in issues if i.get("risk_level") == "高")
    mid = sum(1 for i in issues if i.get("risk_level") == "中")
    low = sum(1 for i in issues if i.get("risk_level") == "低")
    doc.add_paragraph(f"高リスク: {high}件　中リスク: {mid}件　低リスク: {low}件　合計: {len(issues)}件")

    # 詳細
    doc.add_heading("■ 指摘事項詳細", level=2)
    risk_colors = {"高": RGBColor(0xC0, 0x00, 0x00), "中": RGBColor(0xFF, 0x7F, 0x00), "低": RGBColor(0x00, 0x70, 0xC0)}

    for i, issue in enumerate(issues, 1):
        level = issue.get("risk_level", "低")
        p = doc.add_paragraph()
        run = p.add_run(f"【{i}】{issue.get('clause', '')}　（ページ {issue.get('page', '-')}）　リスク: {level}")
        run.bold = True
        run.font.color.rgb = risk_colors.get(level, RGBColor(0, 0, 0))

        doc.add_paragraph(f"問題点: {issue.get('issue', '')}")
        rec = doc.add_paragraph(f"推奨対応: {issue.get('recommendation', '')}")
        rec.runs[0].font.color.rgb = RGBColor(0x00, 0x60, 0x00)
        doc.add_paragraph("")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.get("/history", response_class=HTMLResponse)
async def history_page(request: Request):
    return templates.TemplateResponse("history.html", {"request": request})


@app.get("/api/history")
async def get_history():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    rows = conn.execute("SELECT id, title, created_at, background, focus FROM checks ORDER BY id DESC").fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/api/history/{check_id}")
async def get_history_detail(check_id: int):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    row = conn.execute("SELECT * FROM checks WHERE id = ?", (check_id,)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="履歴が見つかりません")
    data = dict(row)
    data["primary_docs"] = json.loads(data["primary_docs"]) if data["primary_docs"] else []
    data["secondary_docs"] = json.loads(data["secondary_docs"]) if data["secondary_docs"] else []
    data["result"] = json.loads(data["result"]) if data["result"] else []
    return data


@app.delete("/api/history/{check_id}")
async def delete_history(check_id: int):
    conn = sqlite3.connect(DB_PATH)
    conn.execute("DELETE FROM checks WHERE id = ?", (check_id,))
    conn.commit()
    conn.close()
    return {"status": "deleted"}


@app.post("/check")
async def check_contract(request: Request):
    form = await request.form()
    
    background = form.get("background", "")
    focus = form.get("focus", "")
    usage_method = form.get("usage_method", "")
    output_indicator = form.get("output_indicator", "")
    other_context = form.get("other_context", "")
    
    # STEP1: インプット情報（チェック対象）を収集
    primary_docs = []
    for i in range(20):  # 最大20入力
        desc = form.get(f"primary_desc_{i}", "").strip()
        file = form.get(f"primary_file_{i}")
        text = form.get(f"primary_text_{i}", "").strip()
        url = form.get(f"primary_url_{i}", "").strip()
        
        if not desc and not text and not url and (not file or not hasattr(file, 'filename') or not file.filename):
            continue
        
        doc_text = ""
        is_ocr = False
        
        if file and hasattr(file, 'filename') and file.filename:
            filename = file.filename.lower()
            file_bytes = await file.read()
            if filename.endswith(".pdf"):
                doc_text = extract_text_from_pdf(file_bytes)
                is_ocr = True
            elif filename.endswith(".docx"):
                doc_text = extract_text_from_docx(file_bytes)
            elif filename.endswith(".xlsx"):
                doc_text = extract_text_from_xlsx(file_bytes)
            else:
                doc_text = extract_text_from_txt(file_bytes)
        elif text:
            doc_text = text
        elif url:
            doc_text = fetch_text_from_url(url)
        
        if doc_text:
            primary_docs.append({"desc": desc or f"対象資料{i+1}", "text": doc_text, "is_ocr": is_ocr})
    
    # STEP2: インプット付随情報（参照資料）を収集
    secondary_docs = []
    for i in range(20):
        desc = form.get(f"secondary_desc_{i}", "").strip()
        file = form.get(f"secondary_file_{i}")
        text = form.get(f"secondary_text_{i}", "").strip()
        url = form.get(f"secondary_url_{i}", "").strip()
        
        if not desc and not text and not url and (not file or not hasattr(file, 'filename') or not file.filename):
            continue
        
        doc_text = ""
        
        if file and hasattr(file, 'filename') and file.filename:
            filename = file.filename.lower()
            file_bytes = await file.read()
            if filename.endswith(".pdf"):
                doc_text = extract_text_from_pdf(file_bytes)
            elif filename.endswith(".docx"):
                doc_text = extract_text_from_docx(file_bytes)
            elif filename.endswith(".xlsx"):
                doc_text = extract_text_from_xlsx(file_bytes)
            else:
                doc_text = extract_text_from_txt(file_bytes)
        elif text:
            doc_text = text
        elif url:
            doc_text = fetch_text_from_url(url)
        
        if doc_text:
            secondary_docs.append({"desc": desc or f"参照資料{i+1}", "text": doc_text})
    
    if not primary_docs:
        raise HTTPException(status_code=400, detail="チェック対象の契約書を1つ以上入力してください")
    
    try:
        issues = legal_check_with_claude_v2(primary_docs, secondary_docs, background, focus, usage_method, output_indicator, other_context)
    except HTTPException:
        raise
    except Exception as e:
        err = str(e)
        if "ThrottlingException" in err or "Too many tokens" in err or "rate_limit" in err:
            raise HTTPException(status_code=429, detail="AIサービスが混雑しています。しばらく待ってから再試行してください。")
        if "AccessDeniedException" in err or "authentication" in err.lower():
            raise HTTPException(status_code=403, detail="AIモデルへのアクセス権限がありません。APIキーを確認してください。")
        raise HTTPException(status_code=500, detail=f"処理中にエラーが発生しました: {err}")

    # Wordレポート生成
    word_path = generate_word_report(issues, background, focus)

    # 履歴をDBに保存
    try:
        title = primary_docs[0]["desc"] if primary_docs else "無題"
        conn = sqlite3.connect(DB_PATH)
        conn.execute(
            "INSERT INTO checks (title, created_at, background, focus, usage_method, output_indicator, other_context, primary_docs, secondary_docs, result, word_path) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
            (title, datetime.now().isoformat(), background, focus, usage_method, output_indicator, other_context,
             json.dumps([{"desc": d["desc"], "text": d["text"][:500]} for d in primary_docs], ensure_ascii=False),
             json.dumps([{"desc": d["desc"], "text": d["text"][:500]} for d in secondary_docs], ensure_ascii=False),
             json.dumps(issues, ensure_ascii=False), word_path)
        )
        conn.commit()
        conn.close()
    except Exception:
        pass  # 履歴保存失敗してもチェック結果は返す

    return FileResponse(
        word_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="legal_check_report.docx",
        background=None,
    )
